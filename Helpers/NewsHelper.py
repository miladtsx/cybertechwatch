import difflib
import sys
import feedparser
from datetime import datetime, timedelta
from dateutil import parser
from Model import listOfnewsSources, news, localNewsRepository, localRedundantNewsRepository, baselineDate
from multiprocessing.dummy import Pool as ThreadPool 
from docx import Document
from docx.shared import Pt


def DoesItMatch(newItem, localRepoItem):
    if(difflib.SequenceMatcher(None, newItem, localRepoItem).ratio() > 0.9):
        #print '\t[Redundant News] ' + newItem + ' == ' + localRepoItem
        return True
    else:
        return False

# determine the baseline date where News before that day, should be ignored.
# 1 --> just news which was published 24 hours ago
# 3 --> News which published 3 days ago
def getBaselineDate():
    return str(datetime.now() - timedelta(days=3))[:10]


def freshNews(date):
    newsDate = str(parser.parse(date[:-3]))[:10]
    return newsDate >= getBaselineDate()


def UpdateNewsDB():

    # make the Pool of workers
    pool = ThreadPool(len(listOfnewsSources)) 

    # open the urls in their own threads
    # and return the results
    pool.map(GetNewsOnline, listOfnewsSources)

    # close the pool and wait for the work to finish 
    pool.close() 
    pool.join() 

    # for item in listOfnewsSources:
    #     GetNewsOnline(item)


def GetNewsOnline(url):
    try:
        #print '\n[Get] ' + url
        RSSContent = feedparser.parse(url)
        # TODO Check For Null and Errors
        newsGot = RSSContent.entries
        if(len(newsGot) > 0):
            print '\r' + '[' + str(RSSContent.status) + '] ' + url
        else:
            print '\r' + '[' + str(RSSContent.status) + '] ' + url

        for newsItem in newsGot:
            if(not freshNews(newsItem.published)):
                continue
            try:
                newsItemReceived = news(
                    newsItem.title, newsItem.summary, newsItem.link, newsItem.published)
            except:
                newsItemReceived = news(
                    newsItem.title, '', newsItem.link, newsItem.published)

            if len(localNewsRepository) < 1:
                localNewsRepository.append(newsItemReceived)
                continue

            matchFound = False
            for localRep in localNewsRepository:
                if (DoesItMatch(newsItemReceived.title, localRep.title)):
                    matchFound = True
                    localRedundantNewsRepository.append(newsItemReceived)
                    break
                continue

            if(matchFound):
                continue
            else:
                localNewsRepository.append(newsItemReceived)
    except Exception, e:
        print '[Err] ' + url + '\n' + str(e)


def GetNewsOffline():
    NotImplementedError


def PrintNews():
    for i in localNewsRepository:
        print '- ' + i.title + '\n' + i.url + '\n'
    print 'Total News count: ' + \
        str(len(localNewsRepository)) + '\n' + '*' * 100

    for rn in localRedundantNewsRepository:
        print '\t[Redundant News] ' + rn.title + '\n' + rn.url


def saveToFile():
    try:
        if (len(localNewsRepository) < 1):
            print "[Info] Abort Saving report; no News found!"
            return
        localNewsRepository.sort(key=lambda x: x.title)
        
        doc = Document()
        for item in localNewsRepository:
            
            doc.add_heading(str(item.title.encode('utf-8')), level=2)
            p = doc.add_paragraph('')
            p.add_run(str(item.date) + '\n').font.size=Pt(9)
            p.add_run(str(item.url)).font.size=Pt(8)
                                
        doc.save('./reports/' + datetime.today().strftime("%B-%d-%Y.docx"))
        
        print "News File Generated (" + str(len(localNewsRepository)) + ")."
    except Exception,e:
        print '[Err] File operation failed!' + str(e)
