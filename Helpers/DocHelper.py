import docx
from docx import Document
from docx.shared import Pt
import traceback
from Model import  localNewsRepository
from HelperMethods import getFileName

def saveToFile():
    try:
        if (len(localNewsRepository) < 1):
            print ("[Info] Abort Saving report; no News found!")
            return
        localNewsRepository.sort(key=lambda x: x.title)
        
        doc = Document()
        for item in localNewsRepository:
            try:
                doc.add_heading(item.title, level=2)
                p = doc.add_paragraph('')
                p.add_run(str(item.date) + '\n').font.size=Pt(9)
                #p.add_run(str(item.url)).font.size=Pt(8)
                add_hyperlink(p, item.url, str(item.url))


            except Exception as e:
                print ("Exception:" + item.title)
                print ('[Err] ' + '\n' + traceback.print_exc())
                                
        doc.save(getFileName())
        print ("News File Generated (" + str(len(localNewsRepository)) + ").")
    except Exception:
        print (traceback.print_exc())


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    Source: https://github.com/python-openxml/python-docx/issues/74#issuecomment-244602378
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink
